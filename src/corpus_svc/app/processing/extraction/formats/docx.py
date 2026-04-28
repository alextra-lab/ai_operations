"""
DOCX format extractor.
"""

from io import BytesIO
from typing import Any

from shared.logging_utils.fastapi import configure_logging

from ....schemas.chunk import ContentExtractorConfig

logger = configure_logging(service_name="docx_extraction")


async def extract_from_docx(
    content: bytes, config: ContentExtractorConfig, importers: dict[str, Any]
) -> tuple[str, dict]:
    """
    Extract text from DOCX document.

    Args:
        content: DOCX document content as bytes
        config: Content extraction configuration
        importers: Dictionary of lazily imported modules

    Returns:
        Tuple of extracted text and metadata dictionary
    """
    # Lazily import docx
    if "docx" not in importers:
        try:
            import docx

            importers["docx"] = docx
        except ImportError:
            logger.error("python-docx not installed, cannot extract text from DOCX")
            from ..error import ExtractionError

            raise ExtractionError(
                "python-docx not installed, cannot extract text from DOCX. "
                "Install with: pip install python-docx"
            )

    docx_module = importers["docx"]

    metadata = {
        "extraction_method": "python-docx",
        "has_images": False,
    }

    try:
        with BytesIO(content) as docx_stream:
            doc = docx_module.Document(docx_stream)

            # Extract metadata from core properties
            if doc.core_properties:
                if doc.core_properties.author:
                    metadata["author"] = doc.core_properties.author
                if doc.core_properties.title:
                    metadata["title"] = doc.core_properties.title
                if doc.core_properties.created:
                    metadata["created"] = doc.core_properties.created.isoformat()
                if doc.core_properties.modified:
                    metadata["modified"] = doc.core_properties.modified.isoformat()

            # Extract text from paragraphs
            paragraphs = [p.text for p in doc.paragraphs]

            # Extract text from tables
            table_texts = []
            if config.extract_tables and doc.tables:
                for table in doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = "\t".join([cell.text for cell in row.cells])
                        table_text.append(row_text)
                    table_texts.append("\n".join(table_text))

            # Combine all text
            full_text = "\n\n".join(paragraphs)

            # Add tables if extracted
            if table_texts:
                full_text += "\n\n" + "\n\n".join(table_texts)

            # Detect language if configured
            if config.language_detection and full_text:
                from ..text_extractor_utils import detect_language

                metadata["language"] = await detect_language(full_text)

            # Apply length limit if configured
            if config.max_content_length and len(full_text) > config.max_content_length:
                full_text = full_text[: config.max_content_length]
                metadata["truncated"] = True

            return full_text, metadata

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e!s}", exc_info=True)
        from ..error import ExtractionError

        raise ExtractionError(f"Failed to extract text from DOCX: {e!s}")
