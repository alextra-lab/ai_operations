"""
Metadata extractor for document content and file information.
"""

from typing import Any

from shared.logging_utils.fastapi import configure_logging

from ...schemas.document import DocumentType

logger = configure_logging(service_name="metadata_extractor_utils")


class MetadataExtractor:
    """Extract metadata from document content and file information."""

    def __init__(self) -> None:
        """Initialize metadata extractor."""

    async def extract_metadata(
        self,
        content: str | bytes,
        document_type: DocumentType,
        file_name: str,
        file_size: int,
    ) -> dict:
        """
        Extract metadata from document.

        Args:
            content: Document content
            document_type: Type of document
            file_name: Original file name
            file_size: File size in bytes

        Returns:
            Dictionary of metadata
        """
        metadata = {
            "file_name": file_name,
            "file_size": file_size,
            "file_type": document_type,
        }

        # Add file extension
        if "." in file_name:
            metadata["file_extension"] = file_name.split(".")[-1].lower()

        # Document type specific metadata extraction
        if document_type == DocumentType.PDF:
            pdf_metadata = await self._extract_pdf_metadata(content)
            metadata.update(pdf_metadata)
        elif document_type == DocumentType.DOCX:
            docx_metadata = await self._extract_docx_metadata(content)
            metadata.update(docx_metadata)

        return metadata

    async def _extract_pdf_metadata(self, content: str | bytes) -> dict:
        """
        Extract metadata from PDF document.

        Args:
            content: PDF document content

        Returns:
            Dictionary of metadata
        """
        # Convert str to bytes if needed
        content_bytes = content.encode("utf-8") if isinstance(content, str) else content

        metadata: dict[str, Any] = {}

        # Lazily import PyPDF2
        try:
            from io import BytesIO

            from PyPDF2 import PdfReader

            with BytesIO(content_bytes) as pdf_stream:
                reader = PdfReader(pdf_stream)

                # Basic metadata
                metadata["page_count"] = len(reader.pages)

                # Extract document info if available
                if reader.metadata:
                    info = reader.metadata
                    if info.title:
                        metadata["title"] = info.title
                    if info.author:
                        metadata["author"] = info.author
                    if info.subject:
                        metadata["subject"] = info.subject
                    if info.creator:
                        metadata["creator"] = info.creator
                    if info.producer:
                        metadata["producer"] = info.producer
                    if info.creation_date:
                        metadata["creation_date"] = info.creation_date.isoformat()

        except Exception as e:
            logger.warning(f"PDF metadata extraction failed: {e!s}")

        return metadata

    async def _extract_docx_metadata(self, content: str | bytes) -> dict:
        """
        Extract metadata from DOCX document.

        Args:
            content: DOCX document content

        Returns:
            Dictionary of metadata
        """
        # Convert str to bytes if needed
        content_bytes = content.encode("utf-8") if isinstance(content, str) else content

        metadata = {}

        # Lazily import docx
        try:
            from io import BytesIO

            import docx

            with BytesIO(content_bytes) as docx_stream:
                doc = docx.Document(docx_stream)

                # Extract metadata from core properties
                if doc.core_properties:
                    props = doc.core_properties
                    if props.title:
                        metadata["title"] = props.title
                    if props.author:
                        metadata["author"] = props.author
                    if props.comments:
                        metadata["comments"] = props.comments
                    if props.category:
                        metadata["category"] = props.category
                    if props.content_status:
                        metadata["status"] = props.content_status
                    if props.created:
                        metadata["created"] = props.created.isoformat()
                    if props.modified:
                        metadata["modified"] = props.modified.isoformat()
                    if props.last_modified_by:
                        metadata["last_modified_by"] = props.last_modified_by

                # Document statistics
                metadata["paragraph_count"] = len(doc.paragraphs)
                metadata["table_count"] = len(doc.tables)

        except Exception as e:
            logger.warning(f"DOCX metadata extraction failed: {e!s}")

        return metadata
